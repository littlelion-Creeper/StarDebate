"""一辩稿导出工作器 — docx / pdf 生成与 HTML 预览渲染

依赖:
    python-docx, reportlab, PyQt5
"""

from __future__ import annotations

import os
import re

from typing import Optional

# ── docx ───────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    # 提供回退常量，保证模块级 PAGE_SIZE_MAP / ALIGN_MAP_DOCX 等可解析
    class _WDAP:  # WD_ALIGN_PARAGRAPH 的简替代
        LEFT = 0; CENTER = 1; RIGHT = 2; JUSTIFY = 3
    WD_ALIGN_PARAGRAPH = _WDAP
    _DOCX_AVAILABLE = False

# ── reportlab ──────────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4, A5, B5, letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    _PDF_AVAILABLE = True
except ImportError:
    # 提供回退常量（A4 points 值），保证模块级 PAGE_SIZE_MAP 等可解析
    A4 = (595.28, 841.89); A5 = (419.53, 595.28)
    B5 = (498.90, 708.66); letter = (612.00, 792.00)
    cm = 28.35; mm = 2.835
    TA_LEFT = 0; TA_CENTER = 1; TA_RIGHT = 2; TA_JUSTIFY = 3
    _PDF_AVAILABLE = False


# ────────────────────────────────────────────────────────────────────────────
#  常量
# ────────────────────────────────────────────────────────────────────────────

FONT_MAP = {
    "宋体":       ("SimSun",       "宋体"),
    "黑体":       ("SimHei",       "黑体"),
    "楷体":       ("KaiTi",        "楷体"),
    "仿宋":       ("FangSong",     "仿宋"),
    "微软雅黑":   ("Microsoft YaHei", "微软雅黑"),
}

PAGE_SIZE_MAP = {
    "A4":   A4,
    "A5":   A5,
    "B5":   B5,
    "Letter": letter,
}

ALIGN_MAP_CSS = {
    "左对齐": "left",
    "居中对齐": "center",
    "右对齐": "right",
    "两端对齐": "justify",
}

ALIGN_MAP_DOCX = {
    "左对齐":   WD_ALIGN_PARAGRAPH.LEFT,
    "居中对齐": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐":   WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

ALIGN_MAP_PDF = {
    "左对齐":   TA_LEFT,
    "居中对齐": TA_CENTER,
    "右对齐":   TA_RIGHT,
    "两端对齐": TA_JUSTIFY,
}


# ────────────────────────────────────────────────────────────────────────────
#  工具
# ────────────────────────────────────────────────────────────────────────────

def _split_paragraphs(content: str) -> list[str]:
    """按空行 / 单换行切分为段落列表。"""
    raw = re.split(r'\n\s*\n', content.strip())
    return [seg.strip() for seg in raw if seg.strip()]


def _resolve_font_name(font_key: str) -> str:
    """返回 docx 可用的字体名。"""
    entry = FONT_MAP.get(font_key)
    if entry:
        return entry[0]
    return font_key


def _register_chinese_font(font_key: str) -> str | None:
    """尝试注册中文字体到 reportlab，返回注册名或 None。"""
    entry = FONT_MAP.get(font_key)
    if not entry:
        return None
    font_name = entry[0]
    windows_font = os.path.join(
        os.environ.get("WINDIR", "C:\\Windows"), "Fonts",
        f"{font_name}.ttf"
    )
    alt_font = windows_font.replace(".ttf", ".ttc")
    ttf_path = windows_font if os.path.isfile(windows_font) else (
        alt_font if os.path.isfile(alt_font) else None
    )
    if ttf_path:
        try:
            pdfmetrics.registerFont(TTFont(font_name, ttf_path))
            return font_name
        except Exception:
            pass
    return None


def _page_size_mm(page_key: str, orientation: str) -> tuple:
    """返回 (width_mm, height_mm) — 真实毫米值。"""
    base = A4
    if page_key in PAGE_SIZE_MAP:
        base = PAGE_SIZE_MAP[page_key]
    w_pt, h_pt = base
    # reportlab 常量单位是 points (1/72 inch)，需转为 mm
    w_mm = w_pt / mm
    h_mm = h_pt / mm
    if orientation == "横向":
        w_mm, h_mm = h_mm, w_mm
    return w_mm, h_mm


# ────────────────────────────────────────────────────────────────────────────
#  HTML 预览生成
# ────────────────────────────────────────────────────────────────────────────

def generate_preview_html(
    content: str,
    side_label: str = "",
    font_name: str = "宋体",
    font_size: int = 12,
    align: str = "两端对齐",
    indent_chars: int = 2,
    line_spacing: float = 1.5,
    page_size: str = "A4",
    orientation: str = "纵向",
) -> str:
    """生成用于 QWebEngineView 预览的完整 HTML 文档。

    Args:
        content: 一辩稿正文
        side_label: 辩论立场文本（如"支持人工智能发展"），为空时不显示标题头
    """
    paragraphs = _split_paragraphs(content)
    css_align = ALIGN_MAP_CSS.get(align, "justify")
    indent_em = f"{indent_chars}em" if indent_chars > 0 else "0"

    title_html = ""
    if side_label:
        title_html = f'<h2 style="text-align:center; font-weight:bold; font-size:{font_size + 4}pt; text-indent:0; margin-bottom:0.8em;">{_escape_html(side_label)}</h2>'

    p_html = "\n".join(
        f'    <p>{"&nbsp;" * 4 if indent_chars == 0 else ""}{_escape_html(p)}</p>'
        for p in paragraphs
    )

    # 纸张 simulation — 限制 max-width 模拟 A4 宽度
    pw, ph = _page_size_mm(page_size, orientation)
    px_per_mm = 3.78  # ~96dpi
    page_width_px = int(pw * px_per_mm)
    page_height_px = int(ph * px_per_mm)

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    background: #e8eaed;
    display: flex;
    justify-content: center;
    padding: 30px 0;
    font-family: "{font_name}", "SimSun", serif;
    font-size: {font_size}pt;
    line-height: {line_spacing};
    color: #222;
}}
.page {{
    width: {page_width_px}px;
    min-height: {page_height_px}px;
    background: #fff;
    box-shadow: 0 2px 12px rgba(0,0,0,0.12);
    padding: 25mm 30mm;
    text-align: {css_align};
    text-indent: {indent_em};
}}
p {{
    margin: 0 0 0.5em 0;
}}
</style>
</head>
<body>
<div class="page">
{title_html}
{p_html}
</div>
</body>
</html>"""
    return html


def _escape_html(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


# ────────────────────────────────────────────────────────────────────────────
#  DOCX 导出
# ────────────────────────────────────────────────────────────────────────────

def export_to_docx(
    content: str,
    filepath: str,
    side_label: str = "",
    font_name: str = "宋体",
    font_size: int = 12,
    align: str = "两端对齐",
    indent_chars: int = 2,
    line_spacing: float = 1.5,
    page_size: str = "A4",
    orientation: str = "纵向",
) -> str:
    """导出为 .docx 文件，返回路径。

    Args:
        content: 一辩稿正文
        filepath: 保存路径
        side_label: 辩论立场文本（如"支持人工智能发展"），为空时不显示标题头
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError("python-docx 未安装")

    doc = Document()
    resolved_font = _resolve_font_name(font_name)

    # 页面设置
    section = doc.sections[0]
    pw, ph = _page_size_mm(page_size, orientation)
    section.page_width = Mm(pw)
    section.page_height = Mm(ph)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    # ── 居中标题头（显示辩论立场）──
    from docx.oxml.ns import qn
    if side_label:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(12)
        title_p.paragraph_format.first_line_indent = Pt(0)
        title_run = title_p.add_run(side_label)
        title_run.font.name = resolved_font
        title_run.font.size = Pt(font_size + 4)
        title_run.bold = True
        t_rPr = title_run._element.get_or_add_rPr()
        t_rFonts = t_rPr.find(qn('w:rFonts'))
        if t_rFonts is None:
            t_rFonts = title_run._element.makeelement(qn('w:rFonts'), {})
            t_rPr.insert(0, t_rFonts)
        t_rFonts.set(qn('w:eastAsia'), resolved_font)

    paragraphs = _split_paragraphs(content)
    docx_align = ALIGN_MAP_DOCX.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

    for para_text in paragraphs:
        p = doc.add_paragraph()
        p.alignment = docx_align
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = line_spacing
        if indent_chars > 0:
            p.paragraph_format.first_line_indent = Pt(font_size * indent_chars)

        run = p.add_run(para_text)
        run.font.name = resolved_font
        run.font.size = Pt(font_size)

        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = run._element.makeelement(qn('w:rFonts'), {})
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), resolved_font)

    doc.save(filepath)
    return filepath






# ────────────────────────────────────────────────────────────────────────────
#  PDF 导出
# ────────────────────────────────────────────────────────────────────────────

def export_to_pdf(
    content: str,
    filepath: str,
    side_label: str = "",
    font_name: str = "宋体",
    font_size: int = 12,
    align: str = "两端对齐",
    indent_chars: int = 2,
    line_spacing: float = 1.5,
    page_size: str = "A4",
    orientation: str = "纵向",
) -> str:
    """导出为 .pdf 文件，返回路径。

    Args:
        content: 一辩稿正文
        filepath: 保存路径
        side_label: 辩论立场文本（如"支持人工智能发展"），为空时不显示标题头
    """
    if not _PDF_AVAILABLE:
        raise RuntimeError("reportlab 未安装")

    # 注册字体
    registered_name = _register_chinese_font(font_name)
    if not registered_name:
        registered_name = "Helvetica"

    page = PAGE_SIZE_MAP.get(page_size, A4)
    if orientation == "横向":
        page = (page[1], page[0])

    pdf_align = ALIGN_MAP_PDF.get(align, TA_JUSTIFY)
    leading = font_size * line_spacing

    # ── 居中标题头样式（显示辩论立场）──
    title_style = None
    if side_label:
        title_style = ParagraphStyle(
            "SpeechTitle",
            fontName=registered_name,
            fontSize=font_size + 4,
            leading=(font_size + 4) * 1.3,
            alignment=TA_CENTER,
            spaceAfter=14,
        )

    style = ParagraphStyle(
        "SpeechBody",
        fontName=registered_name,
        fontSize=font_size,
        leading=leading,
        alignment=pdf_align,
        leftIndent=0,
        rightIndent=0,
        firstLineIndent=font_size * indent_chars if indent_chars > 0 else 0,
        spaceAfter=6,
    )

    doc = SimpleDocTemplate(
        filepath,
        pagesize=page,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=3.0 * cm,
        rightMargin=3.0 * cm,
    )

    story = []
    # 插入居中标题头（辩论立场）
    if side_label and title_style:
        story.append(Paragraph(side_label, title_style))

    paragraphs = _split_paragraphs(content)
    for para_text in paragraphs:
        safe = para_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        wrapped = safe.replace("\n", "<br/>")
        story.append(Paragraph(wrapped, style))
        story.append(Spacer(1, 2))

    doc.build(story)
    return filepath
